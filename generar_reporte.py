from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

titulo = doc.add_heading('Reporte de Desarrollo — Rudra Kits Peru', level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in titulo.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph(
    'Sitio web de e-commerce y generación de leads para kits de conversión eléctrica\n'
    'Lima, Perú — 2026'
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

GOLD = RGBColor(0xB8, 0x89, 0x0A)

puntos = [
    (
        "1. Definición del producto y mercado objetivo",
        (
            "Se inició definiendo el producto: kits de conversión eléctrica para bicicletas y scooters, "
            "dirigidos al mercado de Lima Metropolitana. El enfoque fue posicionar la marca con un tono "
            "aspiracional y urbano, resaltando velocidad, ahorro y libertad de movilidad como valores centrales."
        )
    ),
    (
        "2. Elección de tecnología: HTML/CSS/JS puro",
        (
            "Se tomó la decisión deliberada de construir el sitio sin frameworks, sin gestor de paquetes "
            "(npm, yarn) ni proceso de build. Todo el sitio corre directamente en el navegador, lo que elimina "
            "complejidad de infraestructura y permite despliegue en cualquier servidor o CDN de forma inmediata. "
            "Solo se requiere un servidor HTTP local (python3 -m http.server) para desarrollo."
        )
    ),
    (
        "3. Estructura de archivos y arquitectura de directorios",
        (
            "Se organizó el proyecto con una arquitectura clara: cada página vive en su propia carpeta "
            "(/scooter/, /beneficios/, /contacto/, /terminos/, /privacidad/) con un index.html dentro para "
            "obtener URLs limpias sin extensión .html. Los assets compartidos se centralizaron en /assets/css/ "
            "y /assets/js/. Un único archivo css/styles.css (963 líneas) concentra todos los estilos globales, "
            "y cada página extiende con su propio CSS específico."
        )
    ),
    (
        "4. Diseño de tokens visuales y sistema de colores",
        (
            "Se definieron los tokens de diseño de la marca: acento principal #F5C518 (amarillo/dorado), "
            "fondo base #0a0a0a (casi negro), y tipografía con espaciado gestionado mediante custom properties CSS. "
            "Esto garantiza coherencia visual en todas las páginas sin depender de un design system externo."
        )
    ),
    (
        "5. Navbar con navegación activa",
        (
            "Se construyó una barra de navegación común a todas las páginas. Incluye el logotipo vectorial SVG "
            "de Rudra (rayo de energía), los 4 enlaces principales, un botón de carrito con ícono SVG y un CTA "
            "primario \"Configurar Kit\". Cada página marca su enlace activo mediante la clase nav-link--active "
            "aplicada directamente en el HTML, sin JavaScript."
        )
    ),
    (
        "6. Hero section con copy de alto impacto",
        (
            "Se desarrolló la sección hero de la página principal con: fondo de imagen con overlay oscuro, "
            "titulares en dos colores (blanco + amarillo), un badge de \"Revoluciona tu movilidad\", descripción "
            "con cifras concretas (velocidades de 30-55 km/h, autonomías de 25-80 km+, 80% de ahorro), dos CTAs "
            "(primario y ghost), fila de estadísticas destacadas e indicador de scroll animado."
        )
    ),
    (
        "7. Sección de catálogo de kits — página de bicicletas",
        (
            "Se diseñó e implementó una grilla de 3 tarjetas de producto: Eco Furia 350W (entrada, 30 km/h, "
            "desde S/1,499), Veloz Titan 500W (el más popular, badge destacado, 40 km/h, desde S/1,899) y "
            "Blitz 1200W (tope de gama, 55 km/h, desde S/2,699). Cada tarjeta muestra ícono SVG, potencia, "
            "velocidad, rango, tiempo de carga, lista de componentes, tabla de upgrades de batería y CTA."
        )
    ),
    (
        "8. Configurador interactivo de kits",
        (
            "Se construyó el configurador, el componente más complejo del sitio. Columna izquierda: selector de "
            "kit con radio buttons estilizados, slider de batería de 3 posiciones con ticks nombrados "
            "(16Ah / 19.2Ah / 25Ah), indicador de capacidad y costo del upgrade, e info box con rango estimado. "
            "Columna derecha (preview en vivo): card con visualización SVG de rueda de motor hub, badge de kit "
            "seleccionado, nombre, especificación, precio total desglosado y CTA \"Agregar al Carrito\"."
        )
    ),
    (
        "9. Lógica JavaScript del configurador (main.js)",
        (
            "Se escribió un IIFE (función autoejecutada) que encapsula toda la lógica del configurador. "
            "El objeto kits contiene todos los datos de los 3 productos: nombre, potencia, velocidad, precio base, "
            "precios de upgrades de batería, rangos y capacidades. Al cambiar el kit o mover el slider, la función "
            "updateUI() recalcula el precio total y actualiza simultáneamente 8 elementos del DOM. El slider se "
            "colorea dinámicamente con un linear-gradient que avanza según la posición seleccionada."
        )
    ),
    (
        "10. Tarjetas de beneficios con efecto flip (CSS 3D)",
        (
            "En la página principal y en /beneficios/, se implementaron tarjetas con efecto de volteo 3D en CSS "
            "puro. Cada tarjeta tiene una cara frontal (ícono + nombre + descripción + CTA) y una cara trasera "
            "(testimonio real con nombre, distrito, kit usado y calificación de 5 estrellas). El JavaScript solo "
            "agrega o quita la clase flipped al hacer click; toda la animación 3D (rotateY) corre en CSS."
        )
    ),
    (
        "11. Página de scooters — catálogo paralelo (/scooter/)",
        (
            "Se replicó y adaptó la estructura de la página principal para scooters, con su propio catálogo de "
            "3 kits: Urban Volt 1500W (entrada para ciudad), Storm Rider 3000W (el más popular) y "
            "Titan Moto 5000W (potencia extrema). La página reutiliza el mismo styles.css global y extiende con "
            "scooter.css para variaciones visuales específicas. scooter.js replica el patrón del configurador "
            "con datos de scooters."
        )
    ),
    (
        "12. Página de beneficios (/beneficios/)",
        (
            "Página dedicada a comunicar el valor de la marca con: hero de alto impacto, sección de números clave "
            "(500+ riders activos, 80% ahorro en combustible, 30 min de instalación, 3+ años de vida útil, "
            "puntuación 4.9 estrellas) y sección de beneficios principales expandida. Todo diseñado para "
            "reforzar la decisión de compra del usuario."
        )
    ),
    (
        "13. Página de contacto (/contacto/) — formulario con validación",
        (
            "Se construyó una página de contacto completa con: hero contextual, sistema de chips para seleccionar "
            "el motivo del contacto, formulario con campos nombre, apellido, email y mensaje. La validación "
            "frontend en contact.js valida cada campo individualmente, muestra mensajes de error inline, limpia "
            "errores al escribir y simula el envío con estado \"Enviando...\" y mensaje de éxito que desaparece "
            "a los 6 segundos."
        )
    ),
    (
        "14. Sección FAQ con acordeón",
        (
            "Dentro de la página de contacto se integró un bloque de preguntas frecuentes con comportamiento de "
            "acordeón. Solo se puede expandir una pregunta a la vez: al abrir una nueva, la anterior se cierra "
            "automáticamente. Usa aria-expanded para accesibilidad."
        )
    ),
    (
        "15. Páginas legales (/terminos/ y /privacidad/)",
        (
            "Se crearon dos páginas de contenido legal con: tabla de contenidos con anclas (#section-id), "
            "legal.js que implementa scroll suave hacia cada sección al hacer click en los enlaces del índice, "
            "y legal.css con tipografía optimizada para lectura de texto largo."
        )
    ),
    (
        "16. Footer completo y consistente",
        (
            "Se diseñó un footer presente en todas las páginas con 4 columnas: Marca (logo, descripción e íconos "
            "de redes sociales: Instagram, Facebook, WhatsApp, YouTube), Kits (links directos a todos los "
            "productos de bicicleta y scooter), Navegación (links a las 6 páginas del sitio) y Contacto "
            "(dirección, teléfono, email y horario). Barra inferior con copyright 2026 y links legales."
        )
    ),
    (
        "17. Estructura de URLs limpias (pretty links)",
        (
            "Se refactorizó el proyecto para que todas las rutas sean limpias y sin extensión .html. Esto se "
            "logró moviendo cada página a su propia carpeta con un index.html dentro. Los assets se centralizaron "
            "en /assets/ y se referencian con rutas absolutas desde la raíz (/assets/css/styles.css), lo que "
            "hace que funcionen correctamente desde cualquier nivel de profundidad de URL."
        )
    ),
    (
        "18. Diseño responsivo y accesibilidad base",
        (
            "Los estilos globales incluyen meta viewport correctamente configurado, etiquetas aria-label en "
            "botones sin texto visible (carrito, redes sociales) y contraste de colores adecuado (texto claro "
            "sobre fondo oscuro). Los SVG inline reemplazan a fuentes de íconos externas, eliminando "
            "dependencias de red."
        )
    ),
    (
        "19. Convenciones de código y mantenibilidad",
        (
            "Se adoptaron convenciones consistentes en todo el proyecto: nomenclatura BEM-inspirada para clases "
            "CSS (config-kit-row--selected, btn-primary, hero-badge), IIFEs en JavaScript para evitar polución "
            "del scope global, atributos data-* para binding entre HTML y JS (data-kit, data-value, data-page), "
            "separación estricta de responsabilidades donde cada JS solo toca su página, y comentarios de "
            "sección con separadores visuales."
        )
    ),
    (
        "20. Control de versiones con Git",
        (
            "El proyecto se gestionó con Git desde el inicio. Se realizaron commits con mensajes descriptivos "
            "que documentan cada etapa del desarrollo (First Upload, feat: added claude.md, feat: pretty links). "
            "El repositorio está alojado en GitHub bajo el usuario francoaalarcon, rama principal main, "
            "permitiendo historial de cambios, reversión y colaboración futura."
        )
    ),
]

for titulo_punto, cuerpo in puntos:
    h = doc.add_heading(titulo_punto, level=2)
    for run in h.runs:
        run.font.color.rgb = GOLD
        run.font.bold = True

    p = doc.add_paragraph(cuerpo)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.2)

doc.add_paragraph()
pie = doc.add_paragraph('Reporte elaborado al 28 de abril de 2026 — Rudra Kits Peru')
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in pie.runs:
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(9)

doc.save('/home/franco/rudrakits/Reporte_Desarrollo_RudraKits.docx')
print("Archivo generado correctamente.")
